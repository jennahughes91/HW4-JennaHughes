#!/usr/bin/env python3
"""
CRM Meeting Report Generator
Usage: python generate_report.py '<json_data>'

JSON shape:
{
  "date": "May 8, 2026",
  "location": "Zoom",
  "participants": [
    {"name": "Sarah Chen", "title": "VP of Sales", "company": "Acme Corp"}
  ],
  "topics": ["Topic one.", "Topic two."],
  "output_path": "/path/to/output.docx"
}
"""

import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
BRAND_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY   = RGBColor(0x59, 0x59, 0x59)
BLACK       = RGBColor(0x00, 0x00, 0x00)
RULE_COLOR  = RGBColor(0xBF, 0xBF, 0xBF)


# ── XML helpers ────────────────────────────────────────────────────────────────
def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_background(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """
    Set borders on a cell. Pass top/bottom/left/right as dicts:
      {"sz": 4, "val": "single", "color": "1F4E79"}
    or as None to clear.
    Inserts tcBorders before shd to maintain correct OOXML element order.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    # OOXML strict uses start/end instead of left/right for tcBorders
    side_map = {"top": "top", "left": "start", "bottom": "bottom", "right": "end"}
    for side in ("top", "left", "bottom", "right"):
        border = kwargs.get(side)
        el = OxmlElement(f"w:{side_map[side]}")
        if border:
            el.set(qn("w:val"),   border.get("val", "single"))
            el.set(qn("w:sz"),    str(border.get("sz", 4)))
            el.set(qn("w:color"), border.get("color", "BFBFBF"))
        else:
            el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    # Insert tcBorders before shd (correct OOXML ordering: tcBorders < shd < tcMar)
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        shd.addprevious(tcBorders)
    else:
        tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    # OOXML strict uses start/end instead of left/right for tcMar
    side_map = [("top", top), ("start", left), ("bottom", bottom), ("end", right)]
    for side, val in side_map:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_paragraph_border_bottom(para, color: RGBColor, sz: int = 8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:color"), hex_color(color))
    bottom.set(qn("w:space"), "4")
    pBdr.append(bottom)
    # pBdr must appear after pStyle (if present) but before everything else
    # Insert after pStyle, otherwise at position 0
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pStyle.addnext(pBdr)
    else:
        pPr.insert(0, pBdr)


def set_paragraph_spacing(para, before: int = 0, after: int = 0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def add_bullet_list_style(para):
    """Apply a simple bullet list style using numPr."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def add_numbering_xml(doc):
    """Inject minimal abstractNum + num XML so bullet list renders correctly."""
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        numbering_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
            'xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex" '
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            'xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink" '
            'xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:oel="http://schemas.microsoft.com/office/2019/extlst" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'xmlns:w10="urn:schemas-microsoft-com:office:word" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
            'xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex" '
            'xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid" '
            'xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml" '
            'xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash" '
            'xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex" '
            'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
            'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
            'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
            'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14">'
            '<w:abstractNum w:abstractNumId="0" w15:restartNumberingAfterBreak="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/>'
            '<w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="&#x2022;"/>'
            '<w:lvlJc w:val="left"/>'
            '<w:pPr>'
            '<w:ind w:left="360" w:hanging="260"/>'
            '</w:pPr>'
            '<w:rPr>'
            '<w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>'
            '</w:rPr>'
            '</w:lvl>'
            '</w:abstractNum>'
            '<w:num w:numId="1">'
            '<w:abstractNumId w:val="0"/>'
            '</w:num>'
            '</w:numbering>'
        )
        numbering_part = doc.part._add_numbering_part()

    # Simpler approach: just inject abstractNum and num elements directly
    numbering_el = doc.part.numbering_part._element

    # Check if abstractNum 0 already exists
    existing = numbering_el.findall(qn("w:abstractNum"))
    if not existing:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), "0")
        mltype = OxmlElement("w:multiLevelType")
        mltype.set(qn("w:val"), "hybridMultilevel")
        abstract.append(mltype)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
        numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "bullet"); lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText"); lvlText.set(qn("w:val"), "•"); lvl.append(lvlText)
        lvlJc = OxmlElement("w:lvlJc"); lvlJc.set(qn("w:val"), "left"); lvl.append(lvlJc)
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "260"); pPr.append(ind)
        lvl.append(pPr)
        abstract.append(lvl)
        numbering_el.insert(0, abstract)

    existing_nums = numbering_el.findall(qn("w:num"))
    if not existing_nums:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstractNumId = OxmlElement("w:abstractNumId")
        abstractNumId.set(qn("w:val"), "0")
        num.append(abstractNumId)
        numbering_el.append(num)


# ── Document builder ──────────────────────────────────────────────────────────
def build_report(date: str, location: str, participants: list, topics: list, output_path: str):
    today = datetime.now().strftime("%B %d, %Y")

    doc = Document()

    # Page size: US Letter, 1" margins
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ── Title ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("CRM Meeting Report")
    title_run.font.name   = "Arial"
    title_run.font.size   = Pt(20)
    title_run.font.bold   = True
    title_run.font.color.rgb = BRAND_BLUE
    set_paragraph_spacing(title_para, before=0, after=80)
    add_paragraph_border_bottom(title_para, BRAND_BLUE, sz=12)

    doc.add_paragraph()  # spacer

    # ── Date / Location ────────────────────────────────────────────────────────
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Inches(1.2), Inches(5.3)]
    for row_idx, (label, value) in enumerate([("Date:", date), ("Location:", location)]):
        row = meta_table.rows[row_idx]
        # Label cell
        lc = row.cells[0]
        lc.width = col_widths[0]
        set_cell_background(lc, WHITE)
        set_cell_border(lc, top=None, bottom=None, left=None, right=None)
        set_cell_margins(lc, top=60, bottom=60, left=0, right=100)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = "Arial"; lr.font.size = Pt(11); lr.font.bold = True
        lr.font.color.rgb = DARK_GRAY

        # Value cell
        vc = row.cells[1]
        vc.width = col_widths[1]
        set_cell_background(vc, WHITE)
        set_cell_border(vc, top=None, bottom=None, left=None, right=None)
        set_cell_margins(vc, top=60, bottom=60, left=0, right=0)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = "Arial"; vr.font.size = Pt(11)
        vr.font.color.rgb = BLACK

    doc.add_paragraph()  # spacer

    # ── Section heading helper ─────────────────────────────────────────────────
    def section_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE
        set_paragraph_spacing(p, before=200, after=120)
        add_paragraph_border_bottom(p, BRAND_BLUE, sz=6)

    # ── Participants table ─────────────────────────────────────────────────────
    section_heading("Meeting Participants")

    # 3 columns: Name (3000), Title (3180), Company (3180) — in DXA units (1440=1in)
    # Full content width = 6.5" = 9360 DXA; split ~= 2.08" / 2.21" / 2.21"
    ptable = doc.add_table(rows=1, cols=3)
    ptable.style = "Table Grid"
    ptable.alignment = WD_TABLE_ALIGNMENT.LEFT

    p_col_widths = [Inches(2.08), Inches(2.21), Inches(2.21)]
    headers = ["Name", "Title", "Company"]

    # Header row
    hrow = ptable.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.width = p_col_widths[i]
        set_cell_background(cell, BRAND_BLUE)
        set_cell_border(cell,
            top={"val": "single", "sz": 4, "color": hex_color(BRAND_BLUE)},
            bottom={"val": "single", "sz": 4, "color": hex_color(BRAND_BLUE)},
            left={"val": "single", "sz": 4, "color": hex_color(BRAND_BLUE)},
            right={"val": "single", "sz": 4, "color": hex_color(BRAND_BLUE)},
        )
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.name = "Arial"; run.font.size = Pt(10); run.font.bold = True
        run.font.color.rgb = WHITE

    # Data rows
    if not participants:
        participants = [{"name": "No participants recorded", "title": "", "company": ""}]

    for row_i, person in enumerate(participants):
        row = ptable.add_row()
        bg = LIGHT_BLUE if row_i % 2 == 1 else WHITE
        for col_i, key in enumerate(["name", "title", "company"]):
            cell = row.cells[col_i]
            cell.width = p_col_widths[col_i]
            set_cell_background(cell, bg)
            set_cell_border(cell,
                top={"val": "single", "sz": 2, "color": hex_color(RULE_COLOR)},
                bottom={"val": "single", "sz": 2, "color": hex_color(RULE_COLOR)},
                left={"val": "single", "sz": 2, "color": hex_color(RULE_COLOR)},
                right={"val": "single", "sz": 2, "color": hex_color(RULE_COLOR)},
            )
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(person.get(key, "") or "")
            run.font.name = "Arial"; run.font.size = Pt(10)
            run.font.color.rgb = BLACK

    doc.add_paragraph()  # spacer

    # ── Topics ────────────────────────────────────────────────────────────────
    section_heading("Key Topics Discussed")

    if topics:
        # Add bullet numbering part
        try:
            doc.part.numbering_part  # check if already exists
        except Exception:
            pass
        add_numbering_xml(doc)

        for topic in topics:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            run = p.add_run(topic)
            run.font.name = "Arial"; run.font.size = Pt(11)
            run.font.color.rgb = BLACK
            set_paragraph_spacing(p, before=60, after=60)
    else:
        p = doc.add_paragraph()
        run = p.add_run("No topics recorded.")
        run.font.name = "Arial"; run.font.size = Pt(11)
        run.font.italic = True; run.font.color.rgb = DARK_GRAY

    # ── Footer line ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f"Prepared by: Claude  |  Report generated: {today}")
    footer_run.font.name = "Arial"; footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = DARK_GRAY
    # Set spacing and pBdr BEFORE alignment so pBdr stays before jc in pPr order
    set_paragraph_spacing(footer_para, before=300, after=0)
    pPr = footer_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "4")
    top_border.set(qn("w:color"), hex_color(RULE_COLOR))
    top_border.set(qn("w:space"), "4")
    pBdr.append(top_border)
    # Insert pBdr before spacing so it sits correctly in pPr element order
    spacing_el = pPr.find(qn("w:spacing"))
    if spacing_el is not None:
        spacing_el.addprevious(pBdr)
    else:
        pPr.insert(0, pBdr)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fix python-docx default: <w:zoom w:val="bestFit"/> is missing required w:percent attribute
    settings_el = doc.settings.element
    for zoom in settings_el.findall(qn("w:zoom")):
        if zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")

    doc.save(output_path)
    print(f"Report saved to: {output_path}")


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_report.py '<json_data>'")
        sys.exit(1)

    raw = sys.argv[1]
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"Invalid JSON: {e}")
        sys.exit(1)

    build_report(
        date        = data.get("date", "Not specified"),
        location    = data.get("location", "Not specified"),
        participants= data.get("participants", []),
        topics      = data.get("topics", []),
        output_path = data["output_path"],
    )
